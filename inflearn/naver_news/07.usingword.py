from docx import Document
# 1. 워드 생성햐기
document = Document()

# 2. 워드 데이터 추가하기
document.add_heading('기사제목', level=0)
document.add_paragraph('기사 링크')
document.add_paragraph('기사 본문')

# 3. 워드 저장하기
document.save("test.docx")