import requests
from bs4 import BeautifulSoup
import pyautogui
import time
from docx import Document

# 사용자 입력
keyword = pyautogui.prompt("검색어를 입력하세요")
lastpage = int(pyautogui.prompt("몇 페이지까지 크롤링 할까요?"))

# 워드 문서 생성하기
document = Document()


page_num = 1
for i in range(1,lastpage * 10, 10):
    print(f"{page_num}페이지 크롤링 중입니다")
    response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}&start={i}")
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')
    articles = soup.select("div.info_group") # 뉴스 기사 div 10개 추출

    for article in articles:
        links = article.select("a.info") # 리스트
        if len(links) >= 2: # 링크가 2개 이상이면
            url = links[1].attrs['href'] # 두번째 링크의 href를 추출
            response = requests.get(url, headers={'User-agent':'Mozila/5.0'})
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            # 만약 스포츠 뉴스라면
            if "entertain" in response.url:
                title = soup.select_one(".end_tit")
                content = soup.select_one("#articeBody")
            elif "sports" in response.url:
                title = soup.select_one("h4.title")
                content = soup.select_one("#newsEndContents")
                # 본문 내용안에 불필요한 div 삭제, p 삭제
                divs = content.select("div")
                for div in divs:
                    div.decompose()
                paragraphs = content.select("p")
                for p in paragraphs:
                    p.decompose()
            else:
                title = soup.select_one("#title")
                content = soup.select_one("#articleBodyContents")

            print("==========링크==========\n", url)
            print("==========제목==========\n", title.text.strip())
            print("==========본문==========\n", content.text.strip())

            # 워드에 제목, 링크, 본문 저장하기
            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())

            time.sleep(0.3)
    page_num = page_num + 1

# 워드 문서 저장하기
document.save(f"{keyword}_result.docx")
